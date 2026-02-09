"""Document filler for populating target documents with matched data."""
import io
import logging
import json
import boto3
import os

logger = logging.getLogger(__name__)

def fill_document(file_bytes, filename, field_mappings):
    """
    Fill target document with matched data.
    
    Args:
        file_bytes: bytes of target document
        filename: name of target file
        field_mappings: dict of {target_field: {'value': value, ...}}
    
    Returns:
        bytes of filled document
    """
    ext = filename.lower().split('.')[-1]
    
    try:
        if ext == 'docx':
            return fill_docx(file_bytes, field_mappings)
        elif ext == 'xlsx':
            return fill_xlsx(file_bytes, field_mappings)
        elif ext == 'txt':
            return fill_txt(file_bytes, field_mappings)
        else:
            raise ValueError(f"Filling not supported for format: {ext}")
    except Exception as e:
        logger.error(f"Failed to fill {filename}: {e}")
        raise

def fill_docx(file_bytes, field_mappings):
    """Fill DOCX document using LLM for intelligent placement."""
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    
    # Extract all text from document
    full_text = "\n".join([para.text for para in doc.paragraphs])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text
    
    # First try simple placeholder replacement
    for para in doc.paragraphs:
        for target_field, data in field_mappings.items():
            value = str(data.get('value', ''))
            placeholders = [
                f"{{{target_field}}}",
                f"[{target_field}]",
                f"__{target_field}__"
            ]
            for placeholder in placeholders:
                if placeholder in para.text:
                    para.text = para.text.replace(placeholder, value)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for target_field, data in field_mappings.items():
                    value = str(data.get('value', ''))
                    placeholders = [
                        f"{{{target_field}}}",
                        f"[{target_field}]",
                        f"__{target_field}__"
                    ]
                    for placeholder in placeholders:
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, value)
    
    # Use LLM to identify where to fill remaining fields
    try:
        bedrock = boto3.client('bedrock-runtime', region_name=os.getenv("AWS_REGION"))
        data_context = "\n".join([f"{field}: {data.get('value', '')}" for field, data in field_mappings.items()])
        
        prompt = f"""Analyze this document and tell me where each data value should be inserted. Return JSON mapping of {{"field_name": "location_description"}}.

Available data:
{data_context}

Document text:
{full_text[:3000]}

Return ONLY valid JSON, no explanation."""
        
        response = bedrock.invoke_model(
            modelId="us.anthropic.claude-sonnet-4-5-20250929-v1:0",
            body=json.dumps({
                "anthropic_version": "bedrock-2023-05-31",
                "max_tokens": 1000,
                "messages": [{"role": "user", "content": prompt}]
            })
        )
        
        result = json.loads(response['body'].read())
        llm_text = result['content'][0]['text']
        
        # Parse suggestions and apply intelligently
        import re
        json_match = re.search(r'\{.*\}', llm_text, re.DOTALL)
        if json_match:
            suggestions = json.loads(json_match.group())
            
            # Apply suggestions by finding matching text patterns
            for field, location_hint in suggestions.items():
                if field in field_mappings:
                    value = str(field_mappings[field].get('value', ''))
                    # Try to find and fill based on location hint
                    for para in doc.paragraphs:
                        if field.lower() in para.text.lower() and not value in para.text:
                            # Append value to paragraph
                            para.text = para.text.rstrip() + " " + value
                            break
    
    except Exception as e:
        logger.error(f"LLM-assisted filling failed: {e}")
    
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

def fill_xlsx(file_bytes, field_mappings):
    """Fill XLSX document by replacing field placeholders."""
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(file_bytes))
    
    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    for target_field, data in field_mappings.items():
                        value = str(data.get('value', ''))
                        placeholders = [
                            f"{{{target_field}}}",
                            f"[{target_field}]",
                            f"__{target_field}__",
                            target_field
                        ]
                        for placeholder in placeholders:
                            if placeholder in cell.value:
                                cell.value = cell.value.replace(placeholder, value)
    
    output = io.BytesIO()
    wb.save(output)
    return output.getvalue()

def fill_txt(file_bytes, field_mappings):
    """Fill TXT document using LLM to intelligently match questions to data."""
    text = file_bytes.decode('utf-8')
    
    # First try simple placeholder replacement
    for target_field, data in field_mappings.items():
        value = str(data.get('value', ''))
        placeholders = [
            f"{{{target_field}}}",
            f"[{target_field}]",
            f"__{target_field}__"
        ]
        for placeholder in placeholders:
            if placeholder in text:
                text = text.replace(placeholder, value)
    
    # Use LLM to intelligently match questions to data
    try:
        bedrock = boto3.client('bedrock-runtime', region_name=os.getenv("AWS_REGION"))
        
        # Format the data more clearly for the LLM
        data_list = []
        for field, data in field_mappings.items():
            value = data.get('value', '')
            data_list.append(f"{field}: {value}")
        data_context = "\n".join(data_list)
        
        prompt = f"""Fill out this form using ONLY the data provided below. You MUST use the available data to answer every question.

AVAILABLE DATA:
{data_context}

FORM TO FILL:
{text}

RULES:
1. Answer EVERY question using the available data
2. Match "yearly" or "last year" questions to yearly data fields
3. Match "monthly" questions to monthly data fields  
4. If a question asks about multiple years and you only have yearly data, state that only 1 year of data is available
5. NEVER say "Information not available" or "No data available" - use the closest matching data
6. Keep the exact same format with "Question:" and "Answer:" labels

Return the COMPLETE filled form:"""
        
        response = bedrock.invoke_model(
            modelId="us.anthropic.claude-sonnet-4-5-20250929-v1:0",
            body=json.dumps({
                "anthropic_version": "bedrock-2023-05-31",
                "max_tokens": 2000,
                "messages": [{"role": "user", "content": prompt}]
            })
        )
        
        result = json.loads(response['body'].read())
        filled_text = result['content'][0]['text']
        
        return filled_text.encode('utf-8')
        
    except Exception as e:
        logger.error(f"LLM filling failed: {e}")
        # Fallback: return original
        return text.encode('utf-8')
